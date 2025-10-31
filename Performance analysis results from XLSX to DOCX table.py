# -*- coding: utf-8 -*-  # STATISTICS OF OBTAINED RESULTS — FINAL GITHUB VERSION
"""
Generate MDPI-style performance table (Word DOCX)
-------------------------------------------------
- Reads Excel file with YOLO benchmark results
- Formats according to MDPI/IEEE precision standards
- Ensures small float numbers are shown numerically (not in scientific notation)
- Exports a Word .docx table ready for publication
"""

import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import subprocess

# === PATHS ===
input_path = r"D:\PUBLIKACJA\WYNIKITORCH\TORCH.xlsx"
output_path = r"D:\PUBLIKACJA\WYNIKITORCH\TORCH.docx"

# === Load data ===
df = pd.read_excel(input_path)

# === Numeric format settings — disable scientific notation ===
pd.options.display.float_format = '{:.10f}'.format  # full precision, prevents truncation

# === Rounding according to publication standards ===
for col in ["fps_mean", "fps_std", "latency_ms_mean", "latency_ms_std"]:
    if col in df.columns:
        df[col] = df[col].round(2)

# mAP50–95/FPS — mean to 4 decimals, std to 5 decimals (no e-05)
if "map50_95_over_fps_mean" in df.columns:
    df["map50_95_over_fps_mean"] = df["map50_95_over_fps_mean"].apply(lambda x: f"{x:.4f}")

if "map50_95_over_fps_std" in df.columns:
    df["map50_95_over_fps_std"] = df["map50_95_over_fps_std"].apply(lambda x: f"{x:.5f}")

# CV% — two decimal places
for col in df.columns:
    if "cv" in col.lower():
        df[col] = df[col].round(2)

# === Create the document ===
doc = Document()
doc.add_heading("Table 1. Performance metrics of YOLO-based models on RTX GPU (FP16 inference)", level=2)

# === Create the table ===
table = doc.add_table(rows=1, cols=len(df.columns))
table.style = "Table Grid"

# Headers
hdr_cells = table.rows[0].cells
for i, col_name in enumerate(df.columns):
    hdr_cells[i].text = col_name.replace("_", " ").capitalize()
    hdr_cells[i].paragraphs[0].runs[0].bold = True
    hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Data rows
for _, row in df.iterrows():
    row_cells = table.add_row().cells
    for j, val in enumerate(row):
        cell = row_cells[j]
        # ensure numeric formatting instead of scientific notation
        if isinstance(val, float):
            if "map50_95_over_fps_std" in df.columns[j]:
                cell.text = f"{val:.5f}"
            elif "map50_95_over_fps_mean" in df.columns[j]:
                cell.text = f"{val:.4f}"
            elif "fps" in df.columns[j] or "latency" in df.columns[j]:
                cell.text = f"{val:.2f}"
            else:
                cell.text = str(val)
        else:
            cell.text = str(val)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# === Table formatting (borders, font) ===
for row in table.rows:
    for cell in row.cells:
        # thin black MDPI-style borders
        cell._tc.get_or_add_tcPr().append(
            parse_xml(r'<w:tcBorders %s>'
                      r'<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                      r'<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                      r'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                      r'<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                      r'</w:tcBorders>' % nsdecls('w'))
        )
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)

# === Table caption/footnote ===
doc.add_paragraph(
    "Abbreviations: CV – coefficient of variation; FPS – frames per second; "
    "Latency – inference time per frame; mAP50:95/FPS – accuracy-to-speed ratio. "
    "Standard deviations below 1×10⁻⁴ were reported as numeric values (not omitted)."
).italic = True

# === Save document ===
doc.save(output_path)
print(f"✅ MDPI table saved to: {output_path}")

# === (Optional) Automatically open in Word ===
try:
    os.startfile(output_path)  # Windows
except Exception:
    try:
        subprocess.run(["open", output_path])  # macOS
    except Exception:
        print("ℹ️ Could not open the file automatically.")
